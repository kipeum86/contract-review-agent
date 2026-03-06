"""Generate a seller-favorable Stock Purchase Agreement (SPA) as DOCX for testing."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Times New Roman"
    h.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        h.font.size = Pt(14)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(10)
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        h.font.size = Pt(12)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
    else:
        h.font.size = Pt(11)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)


def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.italic = italic
    if align:
        p.alignment = align
    return p


def add_centered(text, bold=False, size=None, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


# ══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_centered("STOCK PURCHASE AGREEMENT", bold=True, size=16)
doc.add_paragraph()
add_centered("by and among", size=11)
doc.add_paragraph()
add_centered("DAVID CHEN, SARAH KIM, AND JAMES LIU", bold=True, size=12)
add_centered("as Sellers,", size=11)
doc.add_paragraph()
add_centered("ABC CO., LTD.", bold=True, size=12)
add_centered("as Purchaser,", size=11)
doc.add_paragraph()
add_centered("and", size=11)
doc.add_paragraph()
add_centered("DEF, INC.", bold=True, size=12)
add_centered("(the \"Company\")", size=11)
doc.add_paragraph()
doc.add_paragraph()
add_centered("Dated as of March 7, 2026", size=11)
doc.add_paragraph()
doc.add_paragraph()
add_centered("Prepared by Wilson & Park LLP", italic=True, size=10)
add_centered("Counsel to the Sellers", italic=True, size=10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (simplified)
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("TABLE OF CONTENTS", level=1)
toc_items = [
    ("ARTICLE I", "DEFINITIONS", 3),
    ("ARTICLE II", "PURCHASE AND SALE OF SHARES", 5),
    ("ARTICLE III", "REPRESENTATIONS AND WARRANTIES OF THE SELLERS", 6),
    ("ARTICLE IV", "REPRESENTATIONS AND WARRANTIES OF THE COMPANY", 8),
    ("ARTICLE V", "REPRESENTATIONS AND WARRANTIES OF THE PURCHASER", 10),
    ("ARTICLE VI", "COVENANTS", 13),
    ("ARTICLE VII", "CONDITIONS TO CLOSING", 15),
    ("ARTICLE VIII", "INDEMNIFICATION", 18),
    ("ARTICLE IX", "TERMINATION", 21),
    ("ARTICLE X", "MISCELLANEOUS", 22),
]
for art, title, pg in toc_items:
    p = doc.add_paragraph()
    p.add_run(f"{art}").bold = True
    p.add_run(f"  {title}")
    tab_run = p.add_run(f"\t{pg}")
    tab_run.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# PREAMBLE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("STOCK PURCHASE AGREEMENT", level=1)

add_para(
    'This STOCK PURCHASE AGREEMENT (this "Agreement") is entered into as of March 7, 2026 '
    "(the \"Effective Date\"), by and among:"
)

add_para(
    "(1) David Chen, an individual residing at 1842 Maple Drive, Palo Alto, California 94301 "
    '("Chen");'
)
add_para(
    "(2) Sarah Kim, an individual residing at 2295 Sand Hill Road, Menlo Park, California 94025 "
    '("Kim");'
)
add_para(
    "(3) James Liu, an individual residing at 780 University Avenue, Palo Alto, California 94301 "
    '("Liu," and together with Chen and Kim, each a "Seller" and collectively the "Sellers");'
)
add_para(
    "(4) ABC Co., Ltd., a corporation organized and existing under the laws of the Republic of Korea, "
    "with its principal offices at 123 Teheran-ro, Gangnam-gu, Seoul 06142, Republic of Korea "
    '(the "Purchaser"); and'
)
add_para(
    "(5) DEF, Inc., a Delaware C-Corporation, with its principal offices at 500 Innovation Way, "
    'Wilmington, Delaware 19801 (the "Company").'
)
add_para(
    "Each of the foregoing is referred to herein individually as a \"Party\" and collectively as the \"Parties.\"",
    italic=True,
)

doc.add_paragraph()
doc.add_heading("RECITALS", level=2)
add_para(
    "WHEREAS, the Sellers collectively own shares of common stock of the Company representing "
    "approximately thirty-five percent (35%) of the total issued and outstanding shares of the Company "
    '(the "Shares");'
)
add_para(
    "WHEREAS, the Company is a Series B stage artificial intelligence chip design company incorporated "
    "under the laws of the State of Delaware;"
)
add_para(
    "WHEREAS, the Purchaser desires to purchase from the Sellers, and the Sellers desire to sell to the "
    "Purchaser, the Shares, subject to and upon the terms and conditions set forth herein; and"
)
add_para(
    "WHEREAS, the Board of Directors of the Company has approved the transactions contemplated hereby."
)
add_para(
    "NOW, THEREFORE, in consideration of the mutual covenants and agreements hereinafter set forth and "
    "for other good and valuable consideration, the receipt and sufficiency of which are hereby "
    "acknowledged, the Parties agree as follows:",
    bold=True,
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# ARTICLE I — DEFINITIONS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("ARTICLE I\nDEFINITIONS", level=1)

definitions = [
    (
        '"Affiliate"',
        "means, with respect to any Person, any other Person that directly or indirectly controls, "
        "is controlled by, or is under common control with, such Person.",
    ),
    (
        '"Business Day"',
        "means any day that is not a Saturday, Sunday, or other day on which commercial banks in "
        "Wilmington, Delaware or Seoul, Republic of Korea are authorized or required by law to remain closed.",
    ),
    (
        '"Closing"',
        'means the consummation of the transactions contemplated by this Agreement, as described in Section 2.3.',
    ),
    (
        '"Closing Date"',
        "means the date on which the Closing occurs, which shall be no later than sixty (60) days "
        "following the Effective Date, or such other date as the Sellers may designate in writing to the Purchaser.",
    ),
    (
        '"Company Intellectual Property"',
        "means all Intellectual Property owned by, licensed to, or used by the Company in the conduct "
        "of its business as currently conducted.",
    ),
    (
        '"Encumbrance"',
        "means any lien, pledge, mortgage, deed of trust, security interest, charge, claim, easement, "
        "encroachment, or other similar encumbrance.",
    ),
    (
        '"Governmental Authority"',
        "means any federal, state, provincial, local, municipal, foreign, or other governmental or "
        "quasi-governmental authority of any nature.",
    ),
    (
        '"Knowledge of the Sellers"',
        "or any similar phrase means the actual knowledge of the applicable Seller, without any duty "
        "of inquiry or investigation, after consultation solely with senior management of the Company "
        "who report directly to such Seller in such Seller's capacity as an officer of the Company. "
        "For the avoidance of doubt, Knowledge of the Sellers shall not include constructive knowledge, "
        "imputed knowledge, or knowledge that any Seller could or should have obtained upon reasonable investigation.",
    ),
    (
        '"Law"',
        "means any federal, state, local, municipal, foreign, international, or multinational law, "
        "statute, ordinance, rule, regulation, order, judgment, or decree.",
    ),
    (
        '"Losses"',
        "means any and all losses, damages, liabilities, deficiencies, claims, actions, judgments, "
        "settlements, interest, awards, penalties, fines, costs, or expenses of whatever kind, "
        "including reasonable attorneys' fees.",
    ),
    (
        '"Material Adverse Change"',
        "or \"Material Adverse Effect\" means any event, occurrence, fact, condition, or change that is, "
        "or would reasonably be expected to be, individually or in the aggregate, materially adverse to "
        "the business, results of operations, condition (financial or otherwise), or assets of the Company, "
        "taken as a whole; provided, however, that none of the following shall be deemed to constitute, "
        "and none of the following shall be taken into account in determining whether there has been, "
        "a Material Adverse Change or Material Adverse Effect: (a) any adverse change in general business "
        "or economic conditions; (b) any adverse change in the financial, banking, or securities markets "
        "(including any disruption thereof and any decline in the price of any security or any market index); "
        "(c) any adverse change in conditions generally affecting the semiconductor, artificial intelligence, "
        "or technology industries; (d) any adverse change arising from or relating to any act of terrorism, "
        "war (whether or not declared), armed conflict, sabotage, pandemic, epidemic, or natural disaster; "
        "(e) any adverse change arising from or relating to changes in applicable Law or accounting standards "
        "or the interpretation thereof; (f) any adverse change resulting from the announcement or pendency "
        "of the transactions contemplated by this Agreement; (g) any adverse change resulting from "
        "compliance with the terms of this Agreement; (h) any failure by the Company to meet internal "
        "projections, budgets, plans, or forecasts (provided that the underlying causes of such failure "
        "may be considered); or (i) any adverse change resulting from actions taken at the written request "
        "of the Purchaser.",
    ),
    (
        '"Organizational Documents"',
        "means, with respect to any Person that is an entity, the certificate of incorporation, "
        "articles of organization, bylaws, operating agreement, partnership agreement, or equivalent "
        "organizational documents of such Person.",
    ),
    (
        '"Person"',
        "means any individual, corporation, partnership, limited liability company, joint venture, "
        "trust, unincorporated organization, association, governmental authority, or other entity.",
    ),
    (
        '"Purchase Price"',
        "means Forty Million United States Dollars (US$40,000,000).",
    ),
    (
        '"Sellers\' Proportionate Share"',
        "means, with respect to each Seller, the percentage of the Shares held by such Seller "
        "relative to the total Shares being sold hereunder, as set forth on Schedule A.",
    ),
    (
        '"Shares"',
        "means an aggregate number of shares of common stock, par value $0.001 per share, of the Company, "
        "representing thirty-five percent (35%) of the total issued and outstanding shares of common stock "
        "of the Company as of the Effective Date, as set forth on Schedule A.",
    ),
    (
        '"Transaction Documents"',
        "means this Agreement and all schedules, exhibits, and ancillary agreements attached hereto "
        "or delivered in connection herewith.",
    ),
]

for term, defn in definitions:
    p = doc.add_paragraph()
    run_term = p.add_run(term)
    run_term.bold = True
    p.add_run(f" {defn}")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# ARTICLE II — PURCHASE AND SALE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("ARTICLE II\nPURCHASE AND SALE OF SHARES", level=1)

doc.add_heading("Section 2.1  Purchase and Sale.", level=2)
add_para(
    "Subject to the terms and conditions of this Agreement, at the Closing, the Sellers shall sell, "
    "transfer, assign, and deliver to the Purchaser, and the Purchaser shall purchase, acquire, and "
    "accept from the Sellers, the Shares, free and clear of all Encumbrances (other than restrictions "
    "on transfer under applicable securities laws and the Company's Organizational Documents)."
)

doc.add_heading("Section 2.2  Purchase Price.", level=2)
add_para(
    "(a) The aggregate purchase price for the Shares shall be the Purchase Price. The allocation of "
    "the Purchase Price among the Sellers shall be as set forth on Schedule A."
)
add_para(
    "(b) At the Closing, the Purchaser shall pay the Purchase Price in full by wire transfer of "
    "immediately available funds to the accounts designated by each Seller in writing at least five (5) "
    "Business Days prior to the Closing Date. The entire Purchase Price shall be paid at Closing without "
    "any holdback, escrow, deduction, offset, or withholding of any kind (except as may be required "
    "under applicable tax Law)."
)

doc.add_heading("Section 2.3  Closing.", level=2)
add_para(
    "(a) The Closing shall take place remotely via the electronic exchange of documents and signatures "
    "on the Closing Date, or at such other time, date, and place as the Sellers and the Purchaser may "
    "mutually agree in writing; provided, however, that the Closing shall occur no later than sixty (60) "
    "days following the Effective Date (the \"Outside Date\"). The Sellers shall have the right, in their "
    "sole discretion, to extend the Outside Date by up to an additional thirty (30) days upon written "
    "notice to the Purchaser."
)
add_para(
    "(b) At the Closing, the Sellers shall deliver to the Purchaser: (i) stock certificates representing "
    "the Shares, duly endorsed for transfer or accompanied by duly executed stock powers; and "
    "(ii) such other documents as may be reasonably requested by the Purchaser."
)
add_para(
    "(c) At the Closing, the Purchaser shall deliver to the Sellers: (i) the Purchase Price in accordance "
    "with Section 2.2(b); (ii) a certificate of a duly authorized officer of the Purchaser certifying "
    "that all conditions to the obligations of the Purchaser under Article VII have been satisfied; "
    "(iii) evidence of all required governmental and regulatory approvals, including approvals under "
    "the Korean Foreign Exchange Transactions Act and any other applicable Korean regulatory requirements; "
    "and (iv) such other documents as may be reasonably requested by the Sellers."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# ARTICLE III — SELLERS' REPRESENTATIONS & WARRANTIES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("ARTICLE III\nREPRESENTATIONS AND WARRANTIES OF THE SELLERS", level=1)

add_para(
    "Each Seller, severally and not jointly, represents and warrants to the Purchaser as of the "
    "Effective Date and as of the Closing Date, solely with respect to such Seller, as follows:"
)

doc.add_heading("Section 3.1  Authority.", level=2)
add_para(
    "Such Seller has full legal capacity and authority to execute and deliver this Agreement and to "
    "consummate the transactions contemplated hereby. This Agreement has been duly executed and delivered "
    "by such Seller and constitutes the legal, valid, and binding obligation of such Seller, enforceable "
    "against such Seller in accordance with its terms, subject to applicable bankruptcy, insolvency, and "
    "similar laws affecting creditors' rights generally and to general principles of equity."
)

doc.add_heading("Section 3.2  Title to Shares.", level=2)
add_para(
    "To the Knowledge of the Sellers, such Seller is the record and beneficial owner of such Seller's "
    "Proportionate Share of the Shares, free and clear of all Encumbrances (other than restrictions on "
    "transfer under applicable securities laws and the Company's Organizational Documents). Upon delivery "
    "of the Shares at the Closing, the Purchaser will acquire good and valid title to such Shares, "
    "subject to such restrictions."
)

doc.add_heading("Section 3.3  No Conflicts.", level=2)
add_para(
    "To the Knowledge of the Sellers, the execution, delivery, and performance by such Seller of this "
    "Agreement does not and will not (a) conflict with or violate any Law applicable to such Seller, or "
    "(b) conflict with, result in any breach of, or constitute a default under any material agreement to "
    "which such Seller is a party, in each case, in any material respect."
)

doc.add_heading("Section 3.4  Litigation.", level=2)
add_para(
    "To the Knowledge of the Sellers, there are no Actions pending or threatened against such Seller "
    "that would, individually or in the aggregate, reasonably be expected to materially impair such "
    "Seller's ability to consummate the transactions contemplated by this Agreement."
)

doc.add_heading("Section 3.5  No Brokers.", level=2)
add_para(
    "To the Knowledge of the Sellers, such Seller has not retained any broker, finder, or investment "
    "banker in connection with the transactions contemplated by this Agreement."
)

add_para(
    "EXCEPT FOR THE REPRESENTATIONS AND WARRANTIES EXPRESSLY SET FORTH IN THIS ARTICLE III, NO SELLER "
    "MAKES ANY OTHER REPRESENTATION OR WARRANTY, EXPRESS OR IMPLIED, AT LAW OR IN EQUITY, IN RESPECT "
    "OF SUCH SELLER, THE SHARES, THE COMPANY, OR ANY OF THE COMPANY'S ASSETS, LIABILITIES, OR "
    "OPERATIONS, AND ANY SUCH OTHER REPRESENTATIONS OR WARRANTIES ARE HEREBY EXPRESSLY DISCLAIMED.",
    bold=True,
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# ARTICLE IV — COMPANY REPRESENTATIONS & WARRANTIES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("ARTICLE IV\nREPRESENTATIONS AND WARRANTIES OF THE COMPANY", level=1)

add_para(
    "The Company represents and warrants to the Purchaser as of the Effective Date and as of the "
    "Closing Date as follows, subject in each case to the qualifications and limitations set forth in "
    "the Company Disclosure Schedule delivered concurrently herewith:"
)

doc.add_heading("Section 4.1  Organization and Good Standing.", level=2)
add_para(
    "The Company is a corporation duly organized, validly existing, and in good standing under the "
    "laws of the State of Delaware and has all requisite corporate power and authority to own, lease, "
    "and operate its assets and to carry on its business as currently conducted."
)

doc.add_heading("Section 4.2  Capitalization.", level=2)
add_para(
    "To the Knowledge of the Sellers, the authorized capital stock of the Company consists of shares "
    "as set forth on Schedule B. All issued and outstanding shares have been duly authorized, validly "
    "issued, fully paid, and nonassessable."
)

doc.add_heading("Section 4.3  Financial Statements.", level=2)
add_para(
    "The Company has made available to the Purchaser unaudited financial statements of the Company for "
    "the fiscal year ended December 31, 2025 (the \"Financial Statements\"). To the Knowledge of the "
    "Sellers, the Financial Statements present fairly, in all material respects, the financial position "
    "of the Company as of their respective dates, subject to normal year-end audit adjustments and the "
    "absence of notes thereto."
)

doc.add_heading("Section 4.4  No Undisclosed Liabilities.", level=2)
add_para(
    "To the Knowledge of the Sellers, the Company does not have any liabilities or obligations of any "
    "nature (whether accrued, absolute, contingent, or otherwise) that are, individually or in the "
    "aggregate, material to the Company, except (a) as reflected in or reserved against in the Financial "
    "Statements, (b) liabilities incurred since December 31, 2025, in the ordinary course of business "
    "consistent with past practice, and (c) liabilities arising under this Agreement."
)

doc.add_heading("Section 4.5  Intellectual Property.", level=2)
add_para(
    "To the Knowledge of the Sellers, the Company owns or has adequate rights to use all Company "
    "Intellectual Property material to the conduct of its business as currently conducted. To the "
    "Knowledge of the Sellers, no third party is infringing upon any Company Intellectual Property "
    "owned by the Company in any material respect."
)

doc.add_heading("Section 4.6  Material Contracts.", level=2)
add_para(
    "To the Knowledge of the Sellers, Schedule C sets forth a true and complete list of all material "
    "contracts to which the Company is a party. To the Knowledge of the Sellers, each such material "
    "contract is valid and binding on the Company and, to the Knowledge of the Sellers, each other "
    "party thereto, and is in full force and effect."
)

doc.add_heading("Section 4.7  Compliance with Laws.", level=2)
add_para(
    "To the Knowledge of the Sellers, the Company is in compliance in all material respects with "
    "all Laws applicable to the conduct of its business as currently conducted."
)

doc.add_heading("Section 4.8  Tax Matters.", level=2)
add_para(
    "To the Knowledge of the Sellers, (a) the Company has timely filed all material Tax returns "
    "required to be filed by it, and (b) the Company has timely paid all material Taxes due and payable. "
    "The Sellers make no representation or warranty with respect to the amount of any Tax attribute "
    "(including net operating losses or Tax credits) of the Company that may be available to the Company "
    "or the Purchaser after the Closing."
)

add_para(
    "THE REPRESENTATIONS AND WARRANTIES SET FORTH IN THIS ARTICLE IV ARE MADE SOLELY AS OF THE DATES "
    "SPECIFIED HEREIN AND ARE SUBJECT IN ALL RESPECTS TO THE COMPANY DISCLOSURE SCHEDULE. THE SELLERS "
    "MAKE NO REPRESENTATIONS OR WARRANTIES WITH RESPECT TO THE COMPANY EXCEPT AS EXPRESSLY SET FORTH "
    "IN THIS ARTICLE IV, AND THE PURCHASER ACKNOWLEDGES THAT IT HAS NOT RELIED ON ANY OTHER "
    "REPRESENTATION OR WARRANTY.",
    bold=True,
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# ARTICLE V — PURCHASER'S REPRESENTATIONS & WARRANTIES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("ARTICLE V\nREPRESENTATIONS AND WARRANTIES OF THE PURCHASER", level=1)

add_para(
    "The Purchaser represents and warrants to the Sellers and the Company as of the Effective Date "
    "and as of the Closing Date as follows:"
)

doc.add_heading("Section 5.1  Organization and Good Standing.", level=2)
add_para(
    "The Purchaser is a corporation duly organized, validly existing, and in good standing under the "
    "laws of the Republic of Korea and has all requisite corporate power and authority to execute, "
    "deliver, and perform its obligations under this Agreement and to consummate the transactions "
    "contemplated hereby."
)

doc.add_heading("Section 5.2  Authority.", level=2)
add_para(
    "The execution, delivery, and performance by the Purchaser of this Agreement and the consummation "
    "of the transactions contemplated hereby have been duly authorized by all necessary corporate action "
    "on the part of the Purchaser, including approval by the Purchaser's board of directors and, if "
    "required, its shareholders. This Agreement has been duly executed and delivered by the Purchaser "
    "and constitutes the legal, valid, and binding obligation of the Purchaser, enforceable against the "
    "Purchaser in accordance with its terms."
)

doc.add_heading("Section 5.3  No Conflicts.", level=2)
add_para(
    "The execution, delivery, and performance by the Purchaser of this Agreement and the consummation "
    "of the transactions contemplated hereby do not and will not (a) conflict with or violate the "
    "Organizational Documents of the Purchaser; (b) conflict with, result in any breach of, constitute "
    "a default under, result in the acceleration of, create in any party the right to accelerate, "
    "terminate, modify, or cancel, or require any notice under, any contract, agreement, lease, license, "
    "instrument, or other arrangement to which the Purchaser is a party or by which the Purchaser or "
    "any of its assets is bound; (c) violate any Law applicable to the Purchaser; or (d) result in the "
    "creation or imposition of any Encumbrance upon any of the assets of the Purchaser."
)

doc.add_heading("Section 5.4  Governmental Approvals.", level=2)
add_para(
    "The Purchaser represents and warrants that it has identified all governmental and regulatory "
    "approvals required in connection with the transactions contemplated hereby, including without "
    "limitation approvals under the Korean Foreign Exchange Transactions Act, the Korean Foreign "
    "Investment Promotion Act, and any other applicable Korean, United States federal, state, or "
    "foreign regulatory requirements. The Purchaser shall be solely responsible for obtaining all such "
    "approvals at its own cost and expense prior to the Closing Date."
)

doc.add_heading("Section 5.5  Financial Capability.", level=2)
add_para(
    "The Purchaser has, and at the Closing will have, sufficient immediately available funds to pay "
    "the Purchase Price in full and to perform all of its obligations under this Agreement. The "
    "Purchaser's obligation to consummate the transactions contemplated by this Agreement is not "
    "contingent upon the Purchaser's ability to obtain financing of any kind. The Purchaser has not "
    "incurred, and will not incur, any obligation, commitment, or understanding to pay any fee or "
    "commission to any broker, finder, or intermediary in connection with the transactions contemplated "
    "by this Agreement."
)

doc.add_heading("Section 5.6  Investment Intent.", level=2)
add_para(
    "The Purchaser is acquiring the Shares for investment for its own account and not with a view to, "
    "or for sale in connection with, any distribution thereof in violation of applicable securities laws. "
    "The Purchaser acknowledges that the Shares have not been registered under the Securities Act of "
    "1933, as amended, or any state securities laws, and that the Shares may not be sold, transferred, "
    "or otherwise disposed of without registration thereunder or an applicable exemption therefrom."
)

doc.add_heading("Section 5.7  Investigation; No Other Representations.", level=2)
add_para(
    "The Purchaser acknowledges and agrees that (a) it has been afforded the opportunity to conduct "
    "its own investigation of the Company, the Sellers, and the Shares; (b) in making its decision to "
    "enter into this Agreement and consummate the transactions contemplated hereby, the Purchaser has "
    "relied solely on its own investigation and the representations and warranties of the Sellers and "
    "the Company expressly set forth in Articles III and IV (as qualified by the Company Disclosure "
    "Schedule); (c) neither the Sellers, the Company, nor any of their respective representatives has "
    "made any representation or warranty to the Purchaser, express or implied, other than as expressly "
    "set forth herein; and (d) the Purchaser hereby disclaims reliance on any representation or warranty "
    "not expressly set forth in this Agreement."
)

doc.add_heading("Section 5.8  Sophistication.", level=2)
add_para(
    "The Purchaser is a sophisticated party and has such knowledge and experience in financial and "
    "business matters as to be capable of evaluating the merits and risks of its investment in the "
    "Shares. The Purchaser has been represented by independent counsel and independent financial advisors "
    "in connection with this Agreement."
)

doc.add_heading("Section 5.9  Compliance with Laws.", level=2)
add_para(
    "The Purchaser is in compliance in all respects with all Laws applicable to the Purchaser and its "
    "business, including all Korean corporate governance, securities, foreign exchange, anti-money "
    "laundering, anti-corruption, and tax laws. The Purchaser has not been subject to any governmental "
    "investigation, enforcement action, or sanction within the past five (5) years."
)

doc.add_heading("Section 5.10  No Litigation.", level=2)
add_para(
    "There are no Actions pending or, to the knowledge of the Purchaser, threatened against the "
    "Purchaser or any of its Affiliates that would (a) materially impair the Purchaser's ability to "
    "consummate the transactions contemplated by this Agreement, or (b) have a material adverse effect "
    "on the Purchaser's business, results of operations, or financial condition."
)

doc.add_heading("Section 5.11  Anti-Corruption.", level=2)
add_para(
    "Neither the Purchaser nor any of its Affiliates, directors, officers, employees, or agents has, "
    "directly or indirectly, (a) made, offered, promised, or authorized any payment or transfer of "
    "anything of value to any government official, political party, or any other Person for the purpose "
    "of influencing any act or decision or securing any improper advantage, or (b) violated any "
    "applicable anti-corruption Law, including the U.S. Foreign Corrupt Practices Act, the Korean "
    "Act on Anti-Corruption and the Establishment of the Anti-Corruption and Civil Rights Commission, "
    "and any similar applicable Laws."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# ARTICLE VI — COVENANTS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("ARTICLE VI\nCOVENANTS", level=1)

doc.add_heading("Section 6.1  Conduct of Business.", level=2)
add_para(
    "From the Effective Date until the Closing Date, the Company shall conduct its business in the "
    "ordinary course consistent with past practice. Nothing in this Section 6.1 shall restrict the "
    "Company's ability to (a) issue additional equity securities, (b) incur indebtedness, (c) enter "
    "into related party transactions, (d) modify compensation arrangements for directors, officers, "
    "and employees, (e) declare or pay dividends, or (f) take any action approved by the Company's "
    "Board of Directors, in each case without the consent of the Purchaser."
)

doc.add_heading("Section 6.2  Access to Information.", level=2)
add_para(
    "Following the Closing, the Company shall provide the Purchaser with annual audited financial "
    "statements of the Company within one hundred twenty (120) days after the end of each fiscal year. "
    "The Company shall have no obligation to provide any other financial or operational information to "
    "the Purchaser, including without limitation monthly or quarterly financial statements, management "
    "reports, budgets, projections, or access to the Company's books and records. The Purchaser's "
    "information rights under this Section 6.2 shall terminate upon the earlier of (i) the date on "
    "which the Purchaser holds less than twenty percent (20%) of the outstanding shares of the Company, "
    "or (ii) the fifth (5th) anniversary of the Closing Date."
)

doc.add_heading("Section 6.3  Non-Competition.", level=2)
add_para(
    "Each Seller covenants and agrees that, during the period commencing on the Closing Date and ending "
    "on the date that is six (6) months following the Closing Date (the \"Restricted Period\"), such "
    "Seller shall not, directly or indirectly, within the State of Delaware (the \"Restricted Territory\"), "
    "engage in, own, manage, operate, control, be employed by, participate in, or be connected in any "
    "manner with the ownership, management, operation, or control of any business that directly competes "
    "with the Company's business as conducted on the Closing Date. Notwithstanding the foregoing, nothing "
    "in this Section 6.3 shall prohibit any Seller from (a) owning up to five percent (5%) of the "
    "outstanding securities of any publicly traded company, (b) engaging in any activity outside the "
    "Restricted Territory, or (c) engaging in any academic, research, advisory, or consulting activity."
)

doc.add_heading("Section 6.4  Non-Solicitation.", level=2)
add_para(
    "During the Restricted Period, each Seller shall not, directly or indirectly, within the Restricted "
    "Territory, solicit for employment any employee of the Company; provided, however, that this "
    "restriction shall not apply to (a) general solicitations of employment not specifically directed at "
    "employees of the Company, (b) any employee who has been terminated by the Company, or (c) any "
    "employee who independently contacts such Seller without solicitation."
)

doc.add_heading("Section 6.5  Confidentiality.", level=2)
add_para(
    "The Purchaser shall maintain in strict confidence and shall not disclose to any Person, without "
    "the prior written consent of the Sellers, any Confidential Information of the Sellers or the "
    "Company. \"Confidential Information\" means all non-public information relating to the Company, the "
    "Sellers, or the transactions contemplated by this Agreement. The obligations of confidentiality "
    "under this Section 6.5 shall survive the Closing and continue for a period of five (5) years "
    "following the Closing Date."
)

doc.add_heading("Section 6.6  Public Announcements.", level=2)
add_para(
    "The Purchaser shall not, and shall cause its Affiliates and representatives not to, issue any press "
    "release or make any public statement with respect to this Agreement or the transactions contemplated "
    "hereby without the prior written consent of the Sellers (which consent may be withheld in the "
    "Sellers' sole and absolute discretion). The Sellers may make any public announcement regarding this "
    "Agreement or the transactions contemplated hereby at any time in their sole discretion."
)

doc.add_heading("Section 6.7  Regulatory Approvals.", level=2)
add_para(
    "The Purchaser shall use its best efforts (at its sole cost and expense) to obtain all governmental "
    "and regulatory approvals required in connection with the transactions contemplated by this Agreement, "
    "including without limitation all approvals under Korean foreign exchange and investment regulations. "
    "The Sellers shall have no obligation to assist the Purchaser in obtaining any such approvals. The "
    "failure of the Purchaser to obtain any required regulatory approval shall not relieve the Purchaser "
    "of its obligation to consummate the transactions contemplated hereby, and such failure shall "
    "constitute a breach of this Agreement by the Purchaser."
)

doc.add_heading("Section 6.8  Drag-Along Rights.", level=2)
add_para(
    "If the Sellers (or any two of the three Sellers) elect to sell, transfer, or otherwise dispose of "
    "all or a majority of the Shares then held by the Sellers to a bona fide third party purchaser "
    "(a \"Drag-Along Sale\"), the Sellers shall have the right to require the Purchaser to sell all of "
    "the Purchaser's Shares in such Drag-Along Sale on the same terms and conditions as the Sellers "
    "(the \"Drag-Along Right\"). The Purchaser shall cooperate fully with the Sellers in connection with "
    "any Drag-Along Sale, including executing all documents and instruments reasonably requested by the "
    "Sellers or the third party purchaser. The Purchaser shall have no tag-along rights in connection "
    "with any sale of Shares by any Seller."
)

doc.add_heading("Section 6.9  Transfer Restrictions.", level=2)
add_para(
    "The Purchaser shall not sell, transfer, assign, pledge, or otherwise dispose of any Shares without "
    "the prior written consent of all three Sellers. Any purported transfer in violation of this Section "
    "6.9 shall be null and void. The Sellers shall have a right of first refusal on any proposed transfer "
    "of Shares by the Purchaser, exercisable for a period of sixty (60) days following written notice of "
    "the proposed transfer. The Sellers shall be free to sell, transfer, or otherwise dispose of any "
    "shares of the Company held by the Sellers without restriction and without providing notice to, or "
    "obtaining the consent of, the Purchaser."
)

doc.add_heading("Section 6.10  Governance.", level=2)
add_para(
    "The Purchaser acknowledges and agrees that the acquisition of the Shares hereunder shall not "
    "entitle the Purchaser to (a) any seat on, or observer rights with respect to, the Board of "
    "Directors of the Company, (b) any right to designate, nominate, or appoint any director or officer "
    "of the Company, (c) any consent or veto rights with respect to any action of the Company or its "
    "Board of Directors, including without limitation any merger, acquisition, consolidation, sale of "
    "assets, issuance of equity, incurrence of indebtedness, entry into related party transactions, "
    "adoption or modification of budgets, or changes to the Company's Organizational Documents, or "
    "(d) any right to call or requisition meetings of the Board of Directors or the shareholders of the "
    "Company, except as required by applicable Law."
)

doc.add_heading("Section 6.11  Anti-Dilution.", level=2)
add_para(
    "The Purchaser acknowledges and agrees that the Company may issue additional shares of common stock, "
    "preferred stock, convertible securities, options, warrants, or other equity interests at any time "
    "and from time to time following the Closing, without restriction and without the consent of, or "
    "notice to, the Purchaser, and that any such issuance may dilute the Purchaser's percentage ownership "
    "in the Company. The Purchaser shall have no preemptive rights, anti-dilution protections, or rights "
    "of first refusal in connection with any future issuance of equity securities by the Company."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# ARTICLE VII — CONDITIONS TO CLOSING
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("ARTICLE VII\nCONDITIONS TO CLOSING", level=1)

doc.add_heading("Section 7.1  Conditions to Obligations of the Sellers.", level=2)
add_para(
    "The obligations of the Sellers to consummate the Closing are subject to the satisfaction (or "
    "waiver by the Sellers) of each of the following conditions:"
)
conditions_sellers = [
    "(a) the representations and warranties of the Purchaser set forth in Article V shall be true and "
    "correct in all respects (without giving effect to any materiality or similar qualifier) as of the "
    "Effective Date and as of the Closing Date as though made on and as of such date;",
    "(b) the Purchaser shall have performed and complied with, in all respects, all covenants and "
    "agreements required to be performed or complied with by the Purchaser under this Agreement on "
    "or prior to the Closing Date;",
    "(c) the Purchaser shall have obtained all governmental and regulatory approvals required in "
    "connection with the transactions contemplated hereby, including without limitation all approvals "
    "under Korean foreign exchange and investment regulations, in form and substance satisfactory to "
    "the Sellers;",
    "(d) the Purchaser shall have delivered evidence, in form and substance satisfactory to the "
    "Sellers, that the Purchaser has sufficient immediately available funds to pay the Purchase Price "
    "in full;",
    "(e) the Purchaser shall have obtained approval from the Purchaser's board of directors for the "
    "transactions contemplated by this Agreement, in form and substance satisfactory to the Sellers;",
    "(f) no Action shall have been commenced or threatened against any Seller or the Company by any "
    "Governmental Authority or other Person arising out of or relating to the transactions contemplated "
    "by this Agreement;",
    "(g) no Material Adverse Change shall have occurred with respect to the Purchaser since the "
    "Effective Date; and",
    "(h) the Purchaser shall have delivered all deliverables required under Section 2.3(c).",
]
for c in conditions_sellers:
    add_para(c)

doc.add_heading("Section 7.2  Conditions to Obligations of the Purchaser.", level=2)
add_para(
    "The obligations of the Purchaser to consummate the Closing are subject to the satisfaction (or "
    "waiver by the Purchaser) of each of the following conditions:"
)
conditions_purchaser = [
    "(a) the representations and warranties of the Sellers set forth in Article III that are qualified "
    "by Knowledge of the Sellers shall be true and correct as of the Effective Date and as of the "
    "Closing Date as though made on and as of such date, except where the failure to be true and correct "
    "would not, individually or in the aggregate, reasonably be expected to have a Material Adverse "
    "Effect on the applicable Seller's ability to consummate the transactions contemplated hereby;",
    "(b) the Sellers shall have performed and complied with, in all material respects, all covenants "
    "and agreements required to be performed or complied with by the Sellers under this Agreement on "
    "or prior to the Closing Date; and",
    "(c) the Sellers shall have delivered the deliverables required under Section 2.3(b).",
]
for c in conditions_purchaser:
    add_para(c)

add_para(
    "For the avoidance of doubt, the Purchaser's obligation to consummate the Closing is not subject "
    "to any financing condition, and the Purchaser shall not be permitted to refuse to consummate the "
    "Closing on the basis that the Purchaser has not obtained or cannot obtain sufficient financing.",
    italic=True,
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# ARTICLE VIII — INDEMNIFICATION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("ARTICLE VIII\nINDEMNIFICATION", level=1)

doc.add_heading("Section 8.1  Survival.", level=2)
add_para(
    "The representations and warranties of the Sellers set forth in Article III and of the Company set "
    "forth in Article IV shall survive the Closing for a period of twelve (12) months following the "
    "Closing Date (the \"Survival Period\") and shall thereafter be of no further force or effect. The "
    "representations and warranties of the Purchaser set forth in Article V shall survive the Closing "
    "indefinitely. The covenants and agreements contained herein shall survive the Closing in accordance "
    "with their respective terms."
)

doc.add_heading("Section 8.2  Indemnification by the Sellers.", level=2)
add_para(
    "Subject to the limitations set forth in this Article VIII, each Seller shall, severally and not "
    "jointly (in proportion to such Seller's Proportionate Share), indemnify, defend, and hold harmless "
    "the Purchaser and its Affiliates, directors, officers, employees, and agents (collectively, the "
    "\"Purchaser Indemnitees\") from and against any and all Losses incurred or sustained by any Purchaser "
    "Indemnitee arising out of or resulting from (a) any breach of or inaccuracy in any representation "
    "or warranty of such Seller set forth in Article III or of the Company set forth in Article IV, or "
    "(b) any breach of any covenant or agreement of such Seller contained in this Agreement."
)

doc.add_heading("Section 8.3  Indemnification by the Purchaser.", level=2)
add_para(
    "The Purchaser shall indemnify, defend, and hold harmless the Sellers and their respective "
    "Affiliates, heirs, executors, administrators, successors, and assigns (collectively, the "
    "\"Seller Indemnitees\") from and against any and all Losses incurred or sustained by any Seller "
    "Indemnitee arising out of or resulting from (a) any breach of or inaccuracy in any representation "
    "or warranty of the Purchaser set forth in Article V, (b) any breach of any covenant or agreement "
    "of the Purchaser contained in this Agreement, or (c) any claim by any third party arising out of "
    "or relating to the Purchaser's ownership of the Shares following the Closing."
)

doc.add_heading("Section 8.4  Limitations on Indemnification.", level=2)
add_para(
    "(a) Basket. The Sellers shall not be required to indemnify any Purchaser Indemnitee under "
    "Section 8.2(a) unless and until the aggregate amount of all Losses for which indemnification "
    "is sought exceeds Two Million United States Dollars (US$2,000,000) (the \"Basket\"), and then only "
    "to the extent such Losses exceed the Basket."
)
add_para(
    "(b) Cap. The aggregate liability of the Sellers under Section 8.2 shall not exceed, in the "
    "aggregate, Two Million United States Dollars (US$2,000,000) (the \"Cap\"), which represents five "
    "percent (5%) of the Purchase Price."
)
add_para(
    "(c) No Limitation on Purchaser's Indemnification. Notwithstanding anything to the contrary herein, "
    "the indemnification obligations of the Purchaser under Section 8.3 shall not be subject to any "
    "basket, deductible, cap, or other limitation of any kind. The Purchaser's indemnification "
    "obligations shall be in addition to, and not in limitation of, any other rights or remedies "
    "available to the Seller Indemnitees under applicable Law."
)
add_para(
    "(d) Exclusive Remedy. Subject to Section 8.4(c), from and after the Closing, the indemnification "
    "provisions of this Article VIII shall be the sole and exclusive remedy of the Parties with respect "
    "to any claim arising out of or relating to any breach of any representation, warranty, covenant, "
    "or agreement contained in this Agreement; provided, however, that nothing herein shall limit any "
    "Party's right to seek equitable relief, including specific performance and injunctive relief."
)

doc.add_heading("Section 8.5  Indemnification Procedures.", level=2)
add_para(
    "(a) Notice. Any Party seeking indemnification (the \"Indemnified Party\") shall give prompt written "
    "notice to the indemnifying party (the \"Indemnifying Party\") of any claim for which indemnification "
    "is sought; provided, however, that the failure to give such timely notice shall not relieve the "
    "Indemnifying Party of its indemnification obligations except to the extent the Indemnifying Party "
    "is actually prejudiced by such failure."
)
add_para(
    "(b) Defense of Third Party Claims. The Indemnifying Party shall have the right to assume the "
    "defense of any third party claim subject to indemnification hereunder, at its sole cost and "
    "expense. If the Sellers are the Indemnifying Party, the Sellers shall have the right to settle "
    "any such claim without the consent of the Purchaser, provided that such settlement includes an "
    "unconditional release of the Purchaser Indemnitees from all liability with respect to such claim. "
    "If the Purchaser is the Indemnifying Party, the Purchaser shall not settle any claim without the "
    "prior written consent of the applicable Seller."
)

doc.add_heading("Section 8.6  Mitigation.", level=2)
add_para(
    "The Purchaser shall take all commercially reasonable steps to mitigate any Losses for which it "
    "seeks indemnification hereunder. The Sellers shall have no obligation to indemnify any Purchaser "
    "Indemnitee for any Losses to the extent that such Losses could have been avoided or mitigated by "
    "the exercise of commercially reasonable efforts by the Purchaser."
)

doc.add_heading("Section 8.7  No Escrow.", level=2)
add_para(
    "The Parties expressly agree that no portion of the Purchase Price shall be held in escrow, reserve, "
    "or holdback of any kind in connection with the indemnification obligations of the Sellers. The "
    "entire Purchase Price shall be paid to the Sellers at the Closing in accordance with Section 2.2."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# ARTICLE IX — TERMINATION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("ARTICLE IX\nTERMINATION", level=1)

doc.add_heading("Section 9.1  Termination.", level=2)
add_para("This Agreement may be terminated at any time prior to the Closing:")
term_items = [
    "(a) by mutual written consent of the Sellers and the Purchaser;",
    "(b) by the Sellers, if the Purchaser has breached any representation, warranty, covenant, or "
    "agreement contained in this Agreement, and such breach has not been cured within ten (10) "
    "Business Days following written notice thereof from the Sellers to the Purchaser;",
    "(c) by the Sellers, if any condition to the obligations of the Sellers set forth in Section 7.1 "
    "has not been satisfied (or waived by the Sellers) on or prior to the Outside Date;",
    "(d) by the Sellers, in their sole discretion, at any time prior to the Closing Date, upon "
    "fifteen (15) days' prior written notice to the Purchaser, if the Sellers determine in good "
    "faith that market conditions have changed materially since the Effective Date;",
    "(e) by the Purchaser, if the Sellers have breached any representation or warranty contained in "
    "Article III in any material respect, and such breach has not been cured within thirty (30) "
    "Business Days following written notice thereof from the Purchaser to the Sellers; or",
    "(f) by either Party if the Closing has not occurred on or prior to the Outside Date (as may be "
    "extended by the Sellers pursuant to Section 2.3(a)); provided, however, that the right to "
    "terminate this Agreement under this Section 9.1(f) shall not be available to the Purchaser if "
    "the Purchaser's breach of this Agreement has been the principal cause of or resulted in the "
    "failure of the Closing to occur on or before such date.",
]
for item in term_items:
    add_para(item)

doc.add_heading("Section 9.2  Effect of Termination.", level=2)
add_para(
    "In the event of termination of this Agreement pursuant to Section 9.1, this Agreement shall "
    "become void and of no further force or effect, and no Party shall have any liability to any other "
    "Party hereunder; provided, however, that (a) this Section 9.2, Section 6.5 (Confidentiality), and "
    "Article X (Miscellaneous) shall survive any termination of this Agreement; and (b) nothing herein "
    "shall relieve the Purchaser from liability for any willful or intentional breach of this Agreement "
    "prior to such termination. For the avoidance of doubt, the Sellers shall not be liable for any "
    "damages arising from any termination of this Agreement pursuant to Section 9.1(d)."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# ARTICLE X — MISCELLANEOUS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("ARTICLE X\nMISCELLANEOUS", level=1)

doc.add_heading("Section 10.1  Governing Law.", level=2)
add_para(
    "This Agreement and all disputes or controversies arising out of or relating to this Agreement "
    "shall be governed by and construed in accordance with the laws of the State of Delaware, without "
    "giving effect to any choice or conflict of law provision or rule that would cause the application "
    "of the laws of any jurisdiction other than the State of Delaware."
)

doc.add_heading("Section 10.2  Dispute Resolution.", level=2)
add_para(
    "(a) Any dispute, controversy, or claim arising out of or relating to this Agreement shall be "
    "resolved exclusively by binding arbitration administered by the American Arbitration Association "
    "(\"AAA\") in accordance with its Commercial Arbitration Rules then in effect. The arbitration shall "
    "be conducted in Wilmington, Delaware by a single arbitrator selected in accordance with the AAA "
    "Rules."
)
add_para(
    "(b) The arbitrator shall have no authority to award punitive, exemplary, or consequential damages "
    "against the Sellers. The Purchaser hereby waives any right to seek or recover punitive, exemplary, "
    "or consequential damages of any kind against the Sellers in any proceeding arising out of or "
    "relating to this Agreement."
)
add_para(
    "(c) Each Party shall bear its own costs and expenses (including attorneys' fees) in connection "
    "with any arbitration or other proceeding hereunder; provided, however, that if the Sellers are "
    "the prevailing party, the Purchaser shall reimburse the Sellers for all reasonable costs and "
    "expenses (including attorneys' fees) incurred in connection with such proceeding."
)

doc.add_heading("Section 10.3  Notices.", level=2)
add_para(
    "All notices, requests, demands, and other communications hereunder shall be in writing and shall "
    "be deemed duly given (a) when delivered personally, (b) when sent by confirmed electronic mail "
    "(with a copy sent by internationally recognized overnight courier service), or (c) three (3) "
    "Business Days after being sent by internationally recognized overnight courier service, in each "
    "case to the addresses set forth on the signature pages hereto or to such other address as a Party "
    "may designate by written notice."
)

doc.add_heading("Section 10.4  Entire Agreement.", level=2)
add_para(
    "This Agreement (including the Schedules and Exhibits hereto) constitutes the entire agreement "
    "among the Parties with respect to the subject matter hereof and supersedes all prior agreements, "
    "understandings, representations, and warranties, both written and oral, among the Parties with "
    "respect to the subject matter hereof."
)

doc.add_heading("Section 10.5  Amendment and Waiver.", level=2)
add_para(
    "This Agreement may be amended or modified only by a written instrument signed by the Sellers and "
    "the Purchaser. Any waiver of any right or remedy hereunder must be in writing and signed by the "
    "Party granting such waiver. No waiver shall be effective unless in writing."
)

doc.add_heading("Section 10.6  Severability.", level=2)
add_para(
    "If any provision of this Agreement is held to be invalid, illegal, or unenforceable, such "
    "provision shall be modified to the minimum extent necessary to make it valid, legal, and "
    "enforceable while preserving its original intent, and the remaining provisions shall continue "
    "in full force and effect."
)

doc.add_heading("Section 10.7  Assignment.", level=2)
add_para(
    "The Purchaser may not assign this Agreement or any of its rights or obligations hereunder without "
    "the prior written consent of all three Sellers. Any Seller may freely assign this Agreement or any "
    "of its rights or obligations hereunder without the consent of the Purchaser."
)

doc.add_heading("Section 10.8  Third Party Beneficiaries.", level=2)
add_para(
    "This Agreement is for the sole benefit of the Parties and their respective permitted successors "
    "and assigns and nothing herein, express or implied, is intended to or shall confer upon any other "
    "Person any legal or equitable right, benefit, or remedy of any nature whatsoever under or by reason "
    "of this Agreement."
)

doc.add_heading("Section 10.9  Expenses.", level=2)
add_para(
    "All costs and expenses incurred in connection with the negotiation, preparation, execution, and "
    "performance of this Agreement and the transactions contemplated hereby shall be borne by the "
    "Purchaser, including without limitation the reasonable legal fees and expenses of the Sellers' "
    "counsel."
)

doc.add_heading("Section 10.10  Counterparts.", level=2)
add_para(
    "This Agreement may be executed in one or more counterparts, each of which shall be deemed an "
    "original but all of which together shall constitute one and the same instrument. Delivery of an "
    "executed counterpart of a signature page to this Agreement by electronic transmission shall be "
    "effective as delivery of a manually executed counterpart."
)

doc.add_heading("Section 10.11  Further Assurances.", level=2)
add_para(
    "The Purchaser shall, at its own cost and expense, execute and deliver such additional documents, "
    "instruments, and agreements and take such further actions as may be reasonably requested by the "
    "Sellers to carry out the purposes and intent of this Agreement and to consummate the transactions "
    "contemplated hereby."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SIGNATURE PAGE
# ══════════════════════════════════════════════════════════════════════════
add_centered("[Signature Page Follows]", bold=True, size=11)
doc.add_page_break()

add_para(
    "IN WITNESS WHEREOF, the Parties have executed this Stock Purchase Agreement as of the date "
    "first written above.",
    bold=True,
)

doc.add_paragraph()
add_para("SELLERS:", bold=True)
doc.add_paragraph()
add_para("_______________________________")
add_para("David Chen")
doc.add_paragraph()
add_para("_______________________________")
add_para("Sarah Kim")
doc.add_paragraph()
add_para("_______________________________")
add_para("James Liu")

doc.add_paragraph()
doc.add_paragraph()
add_para("PURCHASER:", bold=True)
doc.add_paragraph()
add_para("ABC CO., LTD.", bold=True)
doc.add_paragraph()
add_para("By: _______________________________")
add_para("Name:")
add_para("Title:")

doc.add_paragraph()
doc.add_paragraph()
add_para("COMPANY (for purposes of Sections 4.1 through 4.8 and Section 6.1 through 6.2 only):", bold=True)
doc.add_paragraph()
add_para("DEF, INC.", bold=True)
doc.add_paragraph()
add_para("By: _______________________________")
add_para("Name:")
add_para("Title:")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SCHEDULES (Placeholder)
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("SCHEDULE A", level=1)
add_para("SELLERS' SHARES AND PROPORTIONATE SHARES", bold=True)
doc.add_paragraph()

# Add a table for Schedule A
table = doc.add_table(rows=5, cols=4)
table.style = "Table Grid"
headers = ["Seller", "Number of Shares", "Proportionate Share", "Allocation of Purchase Price"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ["David Chen", "3,500,000", "40%", "US$16,000,000"],
    ["Sarah Kim", "2,625,000", "30%", "US$12,000,000"],
    ["James Liu", "2,625,000", "30%", "US$12,000,000"],
    ["Total", "8,750,000", "100%", "US$40,000,000"],
]
for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = val

doc.add_page_break()

doc.add_heading("SCHEDULE B", level=1)
add_para("CAPITALIZATION OF THE COMPANY", bold=True)
doc.add_paragraph()
add_para(
    "As of the Effective Date, the authorized capital stock of the Company consists of:"
)
add_para("(a) 100,000,000 shares of Common Stock, par value $0.001 per share, of which 25,000,000 shares "
         "are issued and outstanding; and")
add_para("(b) 20,000,000 shares of Preferred Stock, par value $0.001 per share, of which 8,000,000 shares "
         "of Series A Preferred Stock and 5,000,000 shares of Series B Preferred Stock are issued and outstanding.")
doc.add_paragraph()
add_para(
    "The Company has reserved 5,000,000 shares of Common Stock for issuance under its 2023 Equity "
    "Incentive Plan, of which 2,500,000 shares are subject to outstanding options."
)

doc.add_page_break()

doc.add_heading("SCHEDULE C", level=1)
add_para("MATERIAL CONTRACTS", bold=True)
doc.add_paragraph()
add_para("[To be provided by the Company prior to Closing]", italic=True)

doc.add_page_break()

doc.add_heading("COMPANY DISCLOSURE SCHEDULE", level=1)
doc.add_paragraph()
add_para("[To be delivered by the Company concurrently with the execution of this Agreement]", italic=True)

# ── Save ────────────────────────────────────────────────────────────────
base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
out_dir = os.path.join(base, "contract-review", "uploads")
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "DEF_Inc_SPA_Seller_Draft_2026-03-07.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
