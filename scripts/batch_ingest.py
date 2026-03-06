"""Batch ingestion: Steps 1-3 (detect, fingerprint, normalize) for all DOCX files in inbox/raw."""
import hashlib
import json
import os
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

BASE = Path(__file__).resolve().parent.parent
INBOX_RAW = BASE / "contract-review" / "library" / "inbox" / "raw"
RUNS_DIR = BASE / "contract-review" / "library" / "runs" / "ingestion"
INDEX_FILE = BASE / "contract-review" / "library" / "indexes" / "documents.json"


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def make_doc_id(filename: str) -> str:
    name = Path(filename).stem
    # Remove revision suffix like _2309개정
    name = re.sub(r"_\d{4}개정$", "", name)
    # Transliterate Korean stage/type markers to English
    # Order matters: longer strings first to avoid partial replacement
    replacements = [
        ("상환전환우선주식", "rcps"),
        ("전환우선주식", "convertible-preferred"),
        ("신주인수권부사채", "bond-with-warrants"),
        ("조건부지분전환계약서", "safe-conditional-equity"),
        ("주주간계약서", "sha"),
        ("전환사채", "convertible-bond"),
        ("보통주식", "common"),
        ("투자계약서", "investment"),
        ("초기", "early"),
        ("중기", "mid"),
        ("후기", "late"),
        ("통합형", "integrated"),
        ("분리형", "separated"),
    ]
    for ko, en in replacements:
        name = name.replace(ko, en)
    # Clean up brackets, underscores, dots, spaces
    name = re.sub(r"[\[\]\(\)\.]+", "-", name)
    name = re.sub(r"[_\s]+", "-", name)
    name = re.sub(r"-+", "-", name)
    name = name.strip("-").lower()
    return name


def extract_docx_to_markdown(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    lines = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            para = Paragraph(element, doc)
            text = para.text.strip()
            if not text:
                lines.append("")
                continue

            style_name = para.style.name if para.style else ""

            # Map heading styles
            if "Heading 1" in style_name or "제목 1" in style_name:
                lines.append(f"# {text}")
            elif "Heading 2" in style_name or "제목 2" in style_name:
                lines.append(f"## {text}")
            elif "Heading 3" in style_name or "제목 3" in style_name:
                lines.append(f"### {text}")
            elif "Heading 4" in style_name or "제목 4" in style_name:
                lines.append(f"#### {text}")
            elif "Title" in style_name:
                lines.append(f"# {text}")
            else:
                lines.append(text)

        elif tag == "tbl":
            table = Table(element, doc)
            if table.rows:
                # Header row
                header = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]
                lines.append("| " + " | ".join(header) + " |")
                lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                # Data rows
                for row in table.rows[1:]:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    lines.append("| " + " | ".join(cells) + " |")
                lines.append("")

    return "\n".join(lines)


def extract_plain_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def load_index() -> dict:
    if INDEX_FILE.exists():
        with open(INDEX_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"version": 1, "updated_at": None, "documents": []}


def main():
    now = datetime.now(timezone.utc)
    timestamp = now.strftime("%Y%m%dT%H%M%S")

    # Find all DOCX files
    files = sorted([
        f for f in INBOX_RAW.iterdir()
        if f.suffix.lower() == ".docx" and f.name != ".gitkeep"
    ])

    if not files:
        print("No DOCX files found in inbox/raw/")
        sys.exit(0)

    print(f"Found {len(files)} DOCX files to process.\n")

    index = load_index()
    existing_hashes = {doc["sha256"] for doc in index.get("documents", [])}

    results = []

    for i, fpath in enumerate(files, 1):
        print(f"[{i}/{len(files)}] Processing: {fpath.name}")

        # Step 2: Fingerprint
        file_hash = sha256_file(fpath)
        doc_id = make_doc_id(fpath.name)

        if file_hash in existing_hashes:
            print(f"  SKIP (duplicate hash): {doc_id}")
            results.append({"file": fpath.name, "doc_id": doc_id, "status": "duplicate"})
            continue

        # Create run directory
        run_dir = RUNS_DIR / f"{timestamp}_{doc_id}"
        run_dir.mkdir(parents=True, exist_ok=True)
        (run_dir / "normalized").mkdir(exist_ok=True)
        (run_dir / "structure").mkdir(exist_ok=True)
        (run_dir / "clauses").mkdir(exist_ok=True)
        (run_dir / "quality").mkdir(exist_ok=True)

        # Step 3: Normalize
        try:
            md_content = extract_docx_to_markdown(fpath)
            plain_content = extract_plain_text(fpath)

            md_path = run_dir / "normalized" / "clean.md"
            txt_path = run_dir / "normalized" / "plain.txt"

            with open(md_path, "w", encoding="utf-8") as f:
                f.write(md_content)
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(plain_content)

            word_count = len(plain_content.split())
            line_count = len(md_content.splitlines())

            # Save run record
            run_record = {
                "doc_id": doc_id,
                "source_file": fpath.name,
                "sha256": file_hash,
                "timestamp": now.isoformat(),
                "normalized": True,
                "md_lines": line_count,
                "word_count": word_count,
            }
            with open(run_dir / "run-record.json", "w", encoding="utf-8") as f:
                json.dump(run_record, f, ensure_ascii=False, indent=2)

            print(f"  OK: {doc_id} ({line_count} lines, {word_count} words)")
            results.append({
                "file": fpath.name,
                "doc_id": doc_id,
                "sha256": file_hash,
                "status": "normalized",
                "run_dir": str(run_dir.relative_to(BASE)),
                "md_lines": line_count,
                "word_count": word_count,
            })

        except Exception as e:
            print(f"  FAIL: {e}")
            results.append({"file": fpath.name, "doc_id": doc_id, "status": "failed", "error": str(e)})

    # Summary
    ok = sum(1 for r in results if r["status"] == "normalized")
    dup = sum(1 for r in results if r["status"] == "duplicate")
    fail = sum(1 for r in results if r["status"] == "failed")

    print(f"\n{'='*60}")
    print(f"SUMMARY: {ok} normalized, {dup} duplicates, {fail} failed")
    print(f"{'='*60}")

    # Write batch summary
    summary_path = RUNS_DIR / f"{timestamp}_batch-summary.json"
    with open(summary_path, "w", encoding="utf-8") as f:
        json.dump({"timestamp": now.isoformat(), "results": results}, f, ensure_ascii=False, indent=2)

    print(f"\nBatch summary written to: {summary_path.relative_to(BASE)}")


if __name__ == "__main__":
    main()
